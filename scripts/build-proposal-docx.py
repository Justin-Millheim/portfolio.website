#!/usr/bin/env python3
"""Build the portfolio review proposal as a styled .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x15, 0x11, 0x0D)
EMBER = RGBColor(0xB4, 0x52, 0x1F)   # darkened ember for white-paper legibility
EMBER_BRIGHT = RGBColor(0xE8, 0x6A, 0x33)
BRASS = RGBColor(0x9A, 0x7B, 0x2E)
MUTED = RGBColor(0x6B, 0x63, 0x57)
HEADER_BG = "211A14"
ROW_BG = "FAF7F1"
NOTE_BG = "FAF3EA"

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hex_fill)
    tcPr.append(sh)

def set_cell_text(cell, runs, bold=False, color=None, size=9.5, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opts in runs:
        r = p.add_run(text)
        r.font.size = Pt(opts.get("size", size))
        r.bold = opts.get("bold", bold)
        c = opts.get("color", (RGBColor(0xFF,0xFF,0xFF) if white else color))
        if c: r.font.color.rgb = c

def add_heading(text, level):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    if level == 1:
        r.font.size = Pt(22); r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(2)
    elif level == 2:
        r.font.size = Pt(15); r.font.color.rgb = EMBER
        p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
        pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),"single"); bottom.set(qn("w:sz"),"12"); bottom.set(qn("w:space"),"4"); bottom.set(qn("w:color"),"E86A33")
        pbdr.append(bottom); p._p.get_or_add_pPr().append(pbdr)
    elif level == 3:
        r.font.size = Pt(12.5); r.font.color.rgb = INK
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(2)
    elif level == 4:
        r.font.size = Pt(11); r.font.color.rgb = EMBER
        p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(2)
    return p

def add_para(segments, after=6, italic=False, size=10.5):
    """segments: str or list of (text, {opts})"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if isinstance(segments, str):
        segments = [(segments, {})]
    for text, opts in segments:
        r = p.add_run(text)
        r.font.size = Pt(opts.get("size", size))
        r.bold = opts.get("bold", False)
        r.italic = opts.get("italic", italic)
        if opts.get("color"): r.font.color.rgb = opts["color"]
        if opts.get("mono"): r.font.name = "Consolas"
    return p

def add_bullet(segments, style="List Bullet"):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    if isinstance(segments, str):
        segments = [(segments, {})]
    for text, opts in segments:
        r = p.add_run(text)
        r.font.size = Pt(opts.get("size", 10.5))
        r.bold = opts.get("bold", False)
        r.italic = opts.get("italic", False)
        if opts.get("color"): r.font.color.rgb = opts["color"]
        if opts.get("mono"): r.font.name = "Consolas"
    return p

def add_note(text_segments):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    cell = t.cell(0,0); shade(cell, NOTE_BG)
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    for text, opts in text_segments:
        r = p.add_run(text); r.font.size = Pt(9.5)
        r.bold = opts.get("bold", False); r.italic = opts.get("italic", False)
        if opts.get("color"): r.font.color.rgb = opts["color"]
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade(hdr[i], HEADER_BG)
        set_cell_text(hdr[i], h, bold=True, white=True, size=9.5)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1: shade(cells[i], ROW_BG)
            set_cell_text(cells[i], val, size=9.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

# ---- COVER ----
p = doc.add_paragraph()
r = p.add_run("DESIGN & STRATEGY PROPOSAL  ·  FOR REVIEW")
r.font.size = Pt(9); r.bold = True; r.font.color.rgb = BRASS
p.paragraph_format.space_after = Pt(2)

add_heading("Portfolio Review & Improvement Proposal", 1)
add_para([("justinmillheim.com", {"bold": True}), ("    |    Prepared for Justin Millheim    |    June 16, 2026", {"color": MUTED})], after=1)
add_para([("Benchmark reference: ", {"color": MUTED}), ("Harshika “Hershey” Chadha", {"color": MUTED, "italic": True}), (" — harshikachadha.com", {"color": MUTED})], after=10)

# ---- 0 ----
add_heading("0. How to read this document", 2)
add_para("This is a proposal, not a change log. Nothing here has been built yet. The goal is to give you a clear menu of improvements so you can mark what is in scope, what is out, and what needs more thought. Once you weigh in, I will implement only the approved items.")
add_para([("Each recommendation carries a priority (P1 = highest leverage, P3 = nice-to-have), an effort estimate, and the exact file(s) it touches in your codebase, so the conversation stays concrete. Section 8 is an approval checklist you can edit directly.", {})])

add_note([
  ("A note on the benchmark, in the interest of honesty. ", {"bold": True}),
  ("Hershey’s live site blocks automated access (the page itself, a reader proxy, and the Wayback Machine all returned 403 or were blocked by this environment’s network policy). So my read of her site combines two things: facts I could verify about her positioning and content (AI/tech product leader and consultant, founder of TrailScale Consulting, ex-EY, podcast host, Women-in-Tech advocate, digital nomad across 20+ countries), and the design patterns that clean, polished product-management portfolios in her genre use consistently. Where a point is an inference from the genre rather than a confirmed detail of her exact site, I say so. If you send me four or five screenshots (hero, a case-study top, a mid-scroll moment, the about section), I will tighten the comparison into specifics.", {}),
])

# ---- 1 ----
add_heading("1. Executive summary", 2)
add_para("Your site is already well built. The information architecture is sound, the “Ink & Ember” brand system is distinctive and consistent, the writing has a real voice, and the engineering is clean (Next.js 14, a thoughtful Reveal component, analytics wired through every link). You are not starting from a weak position.")
add_para([("The gap you are sensing between your site and a “polished” portfolio is mostly about ", {}), ("visual storytelling and proof", {"bold": True}), (", not taste or code quality. Three things drive that perception:", {})])
add_bullet([("Your work is told in words, not shown in pictures. ", {"bold": True}), ("Every project tile and case study is text on a card. Polished portfolios lead with images: cover art, screenshots, device mockups, process artifacts. The eye has nothing to land on, so the site reads as a blog rather than a showcase.", {})], style="List Number")
add_bullet([("The proof is missing or marked “to do.” ", {"bold": True}), ("Your Work page has a clean Problem / What I did / Impact structure, but several Impact rows are placeholders. The single most credibility-defining element of a PM portfolio — quantified outcomes — is currently the least finished.", {})], style="List Number")
add_bullet([("The motion is one note. ", {"bold": True}), ("You have a single fade-and-rise on scroll. Sites that feel “alive” layer motion: staggered entrances, image reveals, hover depth, a scroll-progress cue, the occasional pinned section. You noticed this, and you were right to.", {})], style="List Number")
add_para("Fix those three and the polish gap mostly closes. The rest of this document breaks them into specific, file-level changes, plus a set of smaller refinements.")

# ---- 2 ----
add_heading("2. Your site today (what I reviewed)", 2)
add_para([("Reviewed directly from the codebase on branch ", {}), ("claude/relaxed-fermat-ol1b7k", {"mono": True, "size": 9.5}), (". Current structure:", {})])
add_table(
  ["Surface", "What it contains", "Notes"],
  [
    [[("Home", {"bold": True}), ("  (app/(site)/page.tsx)", {"mono": True, "size": 8.5})], "Hero (eyebrow triad, H1 throughline, sub, “Currently” card), “What I’m about” (3 mode cards), “What I’ve been up to” (3 featured tiles), testimonials carousel, closing CTA.", "Strong copy. Entirely text and icons. No photography or product imagery above the fold or in tiles."],
    [[("Work", {"bold": True}), ("  (WorkClient.tsx)", {"mono": True, "size": 8.5})], "Tag-filter chips, then an accordion of roles grouped Professional / Leadership. Each expands to Problem / What I did / Impact.", "Good bones. Impact rows carry a todo flag; some read as unfinished. Logos render as a white chip with a fallback initial."],
    [[("Projects", {"bold": True}), ("  (ProjectsClient.tsx)", {"mono": True, "size": 8.5})], "Shipped / Blog toggle. Shipped is a filterable grid of text tiles; Blog is a dated feed.", "The Project type has no image field, so tiles are title + tags + blurb only."],
    [[("Blog", {"bold": True}), ("  (blog/[slug])", {"mono": True, "size": 8.5})], "MDX posts with real components: figures, galleries, carousels, a lightbox, Suno embeds.", "Your richest visual layer is buried one click into the blog, not surfaced on Home or Projects."],
    [[("About", {"bold": True}), ("  (about/page.tsx)", {"mono": True, "size": 8.5})], "Real headshot, three well-written paragraphs, contact buttons.", "Genuinely good. The voice is the strongest asset on the whole site."],
  ],
  widths=[1.5, 3.0, 2.0],
)
add_heading("Aesthetic system (from globals.css and brand-reference.md)", 4)
add_bullet([("Color: ", {"bold": True}), ("warm dark “Ink & Ember” — canvas #15110D, cream text #F1E8DC, ember #E86A33, brass #C99A52. Ember used as a spark, not a coat. A real, ownable identity.", {})])
add_bullet([("Type: ", {"bold": True}), ("Fraunces (display), Hanken Grotesk (body), JetBrains Mono (labels). A sophisticated pairing.", {})])
add_bullet([("Motion: ", {"bold": True}), ("a film-grain overlay, a sticky blurred nav, hover lifts, a pulsing “Currently” dot, and the Reveal fade-rise. Tasteful but uniform.", {})])
add_para([("Bottom line: the foundation is good. You do not need a redesign. You need to show the work, prove the work, and let it move.", {"italic": True, "color": MUTED})])

# ---- 3 ----
add_heading("3. The benchmark — what polished PM portfolios (like Hershey’s) do well", 2)
add_para([("Synthesized from her verifiable positioning and the conventions of the clean, modern PM-portfolio genre she sits in. Treat the visual-specific points as patterns to adopt, not as claims about her exact pixels.", {"italic": True, "color": MUTED})])
add_heading("What this genre does well", 4)
for seg in [
  [("Visual-first case studies. ", {"bold": True}), ("Each project opens with a cover image or hero mockup, a one-line outcome, and the role/timeframe. You see the work before you read about it.", {})],
  [("Outcomes stated as numbers. ", {"bold": True}), ("“Cut onboarding time 40%,” “grew MAU from X to Y.” Metrics persuade; prose supports.", {})],
  [("Generous whitespace and a calm rhythm. ", {"bold": True}), ("Fewer things per screen, larger type, more air. Restraint reads as confidence.", {})],
  [("Layered scroll motion. ", {"bold": True}), ("Staggered reveals, images that clip or scale in, sticky/pinned sections, a subtle scroll-progress indicator. Movement that rewards scrolling.", {})],
  [("A human, present hero. ", {"bold": True}), ("A strong portrait or signature visual, a crisp value statement, an obvious primary action.", {})],
  [("Credibility signals up high. ", {"bold": True}), ("Recognizable logos (employers, clients, press), a podcast/speaking presence, testimonials. Trust before the scroll.", {})],
  [("One clear primary CTA ", {"bold": True}), ("repeated at natural stopping points, usually “let’s talk” plus a resume.", {})],
]:
  add_bullet(seg)
add_heading("Where this style can fall short (so you can leapfrog, not just copy)", 4)
for seg in [
  [("Template sameness. ", {"bold": True}), ("Many of these sites (often Framer-built) look identical — same fonts, same fades, same layout. Your warm dark Ink & Ember system is a genuine advantage. Do not trade your identity for a generic clean look.", {})],
  [("Style over substance. ", {"bold": True}), ("Beautiful shells with thin proof. You already write with specificity and have real artifacts, which is the harder half.", {})],
  [("Motion that gets in the way. ", {"bold": True}), ("Overdone scroll-jacking and slow reveals frustrate. The goal is polish, not a theme-park ride.", {})],
]:
  add_bullet(seg)
add_para([("Key finding. ", {"bold": True, "color": EMBER}), ("The difference is not that her site is “better designed” in the abstract. It is that her genre shows and proves by default, while yours currently tells. You have the rarer assets (a distinct brand, a real voice, actual built things). You are under-displaying them.", {"italic": True})])

# ---- 4 ----
add_heading("4. The core diagnosis", 2)
add_para([("If we name the single highest-leverage change: ", {}), ("your portfolio describes work that the visitor never sees.", {"bold": True}), (" A hiring manager scanning for ten seconds gets dense, well-written text and no imagery to anchor it. The polish they feel on the benchmark is largely the feeling of seeing the work land, backed by numbers, carried by intentional motion.", {})])
add_para([("So the strategy is not “make it prettier.” It is three moves, in order: ", {}), ("(1) show the work, (2) prove the work, (3) make it move.", {"bold": True}), (" Everything in Section 5 ladders up to one of those.", {})])

# ---- 5 ----
add_heading("5. Recommendations", 2)
add_heading("Priority matrix (the whole menu at a glance)", 3)
add_table(
  ["#", "Recommendation", "Theme", "Pri", "Effort", "Touches"],
  [
    ["A1", "Cover imagery for projects (data model + tiles)", "Show", "P1", "M", "projects.ts, tiles"],
    ["A2", "Visual featured cards on Home", "Show", "P1", "M", "page.tsx, globals.css"],
    ["A3", "Real case-study template (image-led)", "Show", "P2", "L", "new route, content"],
    ["B1", "Fill every Impact row with an outcome", "Prove", "P1", "S", "experience.ts"],
    ["B2", "Metrics / impact band", "Prove", "P1", "M", "page.tsx, globals.css"],
    ["B3", "Logo trust bar", "Prove", "P2", "S", "page.tsx, assets"],
    ["C1", "Staggered reveals + motion variants", "Move", "P1", "S", "Reveal.tsx"],
    ["C2", "Image clip/scale reveals", "Move", "P2", "M", "new RevealImage"],
    ["C3", "Scroll-progress bar + nav-on-scroll", "Move", "P2", "S", "Nav.tsx, layout"],
    ["C4", "One pinned/sticky narrative section", "Move", "P3", "M", "page.tsx / case study"],
    ["D1", "Hero visual + kinetic headline", "Show", "P2", "M", "page.tsx, globals.css"],
    ["E1", "Whitespace + type-scale + rhythm pass", "Polish", "P2", "S", "globals.css"],
    ["E2", "Dial back grain; refine hover/shadow", "Polish", "P3", "S", "globals.css"],
    ["F1", "Resume action in nav + CTA", "Prove", "P2", "S", "Nav.tsx, assets"],
    ["F2", "Work-to-case-study link", "Show", "P3", "S", "WorkClient.tsx"],
    ["G1", "OG image, favicon, metadata, perf", "Polish", "P2", "S", "metadata, public/"],
  ],
  widths=[0.4, 2.6, 0.7, 0.5, 0.6, 1.7],
)
add_para([("Effort: S = under an hour, M = a few hours, L = a half-day or more.", {"size": 9, "color": MUTED})])

detail = [
  ("A. Show the work (visual storytelling)", None),
  ("A1 · Give projects real cover imagery (P1)", "This is the highest-leverage change on the site. content/projects.ts has no image field, so every tile is text. Add an optional cover (and coverAlt) to the Project type and render it at the top of each tile. Why: images give the eye an anchor and instantly read as “portfolio.” Where: content/projects.ts, tiles in page.tsx and ProjectsClient.tsx, .tile styles in globals.css. From you: a cover per project (screenshots, a Notion capture for PMA Command Center, a photo of the engraved gifts, a via-ferrata shot). They need to be real, not fancy. I can also generate on-brand placeholder cards so nothing looks empty during rollout."),
  ("A2 · Turn Home featured tiles into visual cards (P1)", "Once A1 exists, the three featured tiles on Home become image-topped cards: cover, then tags, title, blurb, arrow. This is the first thing after the hero, so it carries much of the polish impression."),
  ("A3 · A true case-study template (P2)", "Today “case studies” are blog posts. A dedicated layout (cover hero, role/timeframe/team, problem, process with images, outcome with metrics, what you’d do differently) is what reads as senior PM work. New route or MDX layout variant. The larger investment; worth it for two or three flagship projects (Adobe Claude Skills, PMA Command Center)."),
  ("B. Prove the work (credibility & metrics)", None),
  ("B1 · Finish the Impact rows (P1)", "In content/experience.ts, several roles carry a todo flag on Impact and render italic. For a PM portfolio this is the most important text on the site. Every role should end on a concrete outcome: a number, an adoption figure, a before/after, a shipped artifact. Where no hard metric exists, a specific qualitative result beats a placeholder. Needs your input — I can draft from what you tell me, but the facts are yours."),
  ("B2 · A metrics / impact band (P1)", "A compact band of three or four headline numbers, on Home (after the hero or after testimonials) and at the top of case studies. Big Fraunces numerals in ember, mono caption beneath. Even “20+ AI skills shipped,” “3 student orgs led” create instant scale."),
  ("B3 · A logo trust bar (P2)", "A quiet row of marks — Adobe, University of Utah, prior employers, PMA / Lassonde — high on Home. You already render company logos on Work, so the pattern partly exists. Keep it monochrome so it stays on-brand and does not fight the ember."),
  ("C. Make it move (motion & scroll)", None),
  ("C1 · Staggered reveals + a small variant library (P1)", "Reveal.tsx is solid and already respects reduced-motion. Extend it so groups of items animate in with a stagger (Framer Motion staggerChildren) rather than all at once. Add two or three named variants (fade-up, fade-in, scale-in) for consistent, intentional motion sitewide. Small change, big “alive” payoff."),
  ("C2 · Image reveals (P2)", "A RevealImage that clips or scales media into view as it enters the viewport. Apply to project covers and blog figures. The signature “polished” motion in the benchmark genre; pairs naturally with A1/A2."),
  ("C3 · Scroll-progress + nav-on-scroll (P2)", "A thin ember scroll-progress bar at the top, and a nav that subtly condenses or sharpens its border after a little scroll. Cheap, reads as considered. Touches Nav.tsx and the layout."),
  ("C4 · One pinned narrative moment (P3)", "A single sticky/pinned section where copy or an image holds while related content advances. Used once it feels premium; everywhere it annoys. Candidate: the “What I’m about” triad, or the top of a flagship case study. Optional flourish."),
  ("D. First impression (hero)", None),
  ("D1 · A hero with a visual, and a more kinetic headline (P2)", "Your hero is all text. Options by ambition: (a) add a portrait or signature visual beside the copy on desktop; (b) animate the H1 per line or per word instead of one block fade; (c) keep the excellent “Currently” card (a genuinely distinctive element most portfolios lack). I recommend (a)+(b): keep the words, give them a face and motion. Touches page.tsx and .hero in globals.css."),
  ("E. Polish (rhythm & finish)", None),
  ("E1 · More air, tighter scale (P2)", "Increase section spacing a notch, widen line-height on large headings slightly, standardize vertical rhythm with a small spacing scale. The benchmark’s calm comes mostly from whitespace. Pure globals.css work."),
  ("E2 · Refine the grain and hovers (P3)", "The film-grain overlay sits at opacity .2 with mix-blend-mode overlay; on some screens it adds visible noise. Test it at half strength, and soften card hover shadows so the lift feels more premium. Subtle."),
  ("F. Navigation & conversion", None),
  ("F1 · A Resume action (P2)", "For a job search, “Resume” is a top-three action and is currently absent from the nav. Add it beside “Connect” and in the closing CTA, linking to a PDF in public/."),
  ("F2 · Link Work to its case study (P3)", "When a role has a matching deep-dive (e.g. Adobe), surface a “Read the case study” link inside the expanded accordion body in WorkClient.tsx, so Work and Projects reinforce each other."),
  ("G. Technical polish", None),
  ("G1 · Metadata, OG image, favicon, performance (P2)", "Ensure every page has a proper title/description, a branded Open Graph image (so shared links look sharp), a real favicon, and that images ship through next/image with correct sizing. Quiet credibility that recruiters notice in link previews."),
]
for title, body in detail:
  if body is None:
    add_heading(title, 3)
  else:
    add_heading(title, 4)
    add_para(body)

# ---- 6 ----
add_heading("6. Suggested sequencing", 2)
add_table(
  ["Phase", "Goal", "Items", "Rough effort"],
  [
    [[("Phase 1 — Quick wins", {"bold": True})], "Make it move and prove, cheaply", "C1, B1, B2, F1, E1", "~1 day"],
    [[("Phase 2 — Show the work", {"bold": True})], "Visual storytelling", "A1, A2, C2, B3", "~1–2 days + assets"],
    [[("Phase 3 — Flagship depth", {"bold": True})], "Senior-PM signal", "A3, D1, C3, F2, G1", "~2–3 days"],
    [[("Phase 4 — Flourish", {"bold": True})], "Optional premium touches", "C4, E2", "~half day"],
  ],
  widths=[1.6, 2.0, 1.6, 1.3],
)
add_para([("If you only do one phase, do Phase 1: cheapest path to the biggest perceived jump, and it needs almost nothing from you except the impact numbers.", {})])

# ---- 7 ----
add_heading("7. What I need from you", 2)
add_bullet([("Impact numbers ", {"bold": True}), ("for the Work roles (the B1 facts). The one thing I cannot invent.", {})])
add_bullet([("Project cover assets ", {"bold": True}), ("(A1), or a green light to generate on-brand placeholders for now.", {})])
add_bullet([("A resume PDF ", {"bold": True}), ("if you want F1.", {})])
add_bullet([("Optional: ", {"bold": True}), ("a few screenshots of Hershey’s site to sharpen the benchmark beyond genre patterns.", {})])
add_bullet([("Brand guardrail confirmation: ", {"bold": True}), ("I am assuming we keep Ink & Ember, dark-only, Fraunces/Hanken/Mono, and the “no logo/monogram” rule. Say so if any of that is open for change.", {})])

# ---- 8 ----
add_heading("8. Approval checklist", 2)
add_para([("Mark each: ", {}), ("YES", {"bold": True}), (" (do it), ", {}), ("NO", {"bold": True}), (" (skip), or ", {}), ("MAYBE", {"bold": True}), (" (discuss). Edit freely.", {})])
add_table(
  ["#", "Item", "Pri", "Decision"],
  [
    ["A1", "Cover imagery for projects", "P1", ""],
    ["A2", "Visual featured cards on Home", "P1", ""],
    ["A3", "Dedicated case-study template", "P2", ""],
    ["B1", "Finish all Impact rows with outcomes", "P1", ""],
    ["B2", "Metrics / impact band", "P1", ""],
    ["B3", "Logo trust bar", "P2", ""],
    ["C1", "Staggered reveals + variant library", "P1", ""],
    ["C2", "Image clip/scale reveals", "P2", ""],
    ["C3", "Scroll-progress bar + nav-on-scroll", "P2", ""],
    ["C4", "One pinned narrative section", "P3", ""],
    ["D1", "Hero visual + kinetic headline", "P2", ""],
    ["E1", "Whitespace + type-scale pass", "P2", ""],
    ["E2", "Grain + hover/shadow refinement", "P3", ""],
    ["F1", "Resume action in nav + CTA", "P2", ""],
    ["F2", "Work-to-case-study link", "P3", ""],
    ["G1", "OG image, favicon, metadata, perf", "P2", ""],
  ],
  widths=[0.5, 4.0, 0.7, 1.5],
)
add_para([("Reply with your decisions (even just “Phase 1 plus A1, A2”) and I will implement only the approved items on claude/relaxed-fermat-ol1b7k, then push for your review. No code will change before you say go.", {"italic": True, "color": MUTED})])

doc.save("/home/user/portfolio.website/portfolio-review-proposal.docx")
print("Saved portfolio-review-proposal.docx")
